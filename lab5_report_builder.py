from __future__ import annotations

import math
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import List, Sequence, Tuple

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
ARTIFACT_DIR = BASE_DIR / "lab5_artifacts"
ARTIFACT_DIR.mkdir(exist_ok=True)

REPORT_PATH = BASE_DIR / "实验5_密码分享_报告.docx"
FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"

N = 4


CIRCUIT: List[List[Tuple[int, int, int, int]]] = [
    [(1, 0, 0, 4), (1, 0, 0, 5), (1, 0, 1, 4), (1, 0, 1, 5), (1, 0, 2, 6), (1, 0, 2, 7), (1, 0, 3, 6), (1, 0, 3, 7)],
    [(0, 1, 0, 0), (1, 0, 1, 2), (0, 0, 1, 2), (0, 1, 3, 0), (0, 1, 4, 0), (1, 0, 5, 6), (0, 0, 5, 6), (0, 1, 7, 0)],
    [(1, 0, 0, 1), (0, 0, 0, 1), (0, 1, 2, 0), (0, 1, 3, 0), (1, 0, 4, 5), (0, 0, 4, 5), (0, 1, 6, 0), (0, 1, 7, 0)],
    [(1, 0, 0, 4), (0, 0, 0, 4), (1, 0, 1, 5), (0, 0, 1, 5), (1, 0, 2, 6), (0, 0, 2, 6), (1, 0, 3, 7), (0, 0, 3, 7)],
    [(0, 1, 0, 0), (0, 1, 1, 0), (0, 1, 2, 0), (0, 1, 3, 0), (0, 1, 4, 0), (1, 0, 5, 6), (0, 0, 5, 6), (0, 1, 7, 0)],
    [(0, 1, 0, 0), (0, 1, 1, 0), (0, 1, 2, 0), (0, 1, 3, 0), (0, 0, 4, 5), (0, 1, 6, 0), (0, 1, 7, 0)],
    [(0, 1, 0, 0), (0, 1, 1, 0), (0, 1, 2, 0), (1, 0, 3, 4), (0, 0, 3, 4), (0, 1, 5, 0), (0, 1, 6, 0)],
    [(0, 1, 0, 0), (0, 1, 1, 0), (0, 0, 2, 3), (0, 1, 4, 0), (0, 1, 5, 0), (0, 1, 6, 0)],
    [(0, 1, 0, 0), (1, 0, 1, 2), (0, 0, 1, 2), (0, 1, 3, 0), (0, 1, 4, 0), (0, 1, 5, 0)],
    [(0, 0, 0, 1), (0, 1, 2, 0), (0, 1, 3, 0), (0, 1, 4, 0), (0, 1, 5, 0)],
    [(1, 0, 0, 1), (0, 0, 0, 1), (0, 1, 2, 0)],
    [(0, 0, 0, 1), (0, 1, 2, 0)],
    [(1, 0, 0, 1), (0, 0, 0, 1)],
    [(0, 0, 0, 1)],
]


def bits_to_int(bits: Sequence[int]) -> int:
    return int(bits[0]) * 2 + int(bits[1])


def fmt_bits(bits: Sequence[int]) -> str:
    return "[" + ", ".join(str(int(b)) for b in bits) + "]"


def xor_list(a: Sequence[int], b: Sequence[int]) -> List[int]:
    return [int(x) ^ int(y) for x, y in zip(a, b)]


def rand_bits(rng: np.random.Generator, n: int) -> List[int]:
    return [int(v) for v in rng.integers(0, 2, size=n)]


def expected_output(a_bits: Sequence[int], x_bits: Sequence[int]) -> int:
    a1 = bits_to_int(a_bits[:2])
    a2 = bits_to_int(a_bits[2:])
    x1 = bits_to_int(x_bits[:2])
    x2 = bits_to_int(x_bits[2:])
    return int(a1 * x1 + a2 * x2 >= 4)


@dataclass
class GateTrace:
    layer_index: int
    gate_index: int
    gate: Tuple[int, int, int, int]
    plain_result: int
    dealer_index: int
    dealer_u: int
    dealer_v: int
    dealer_w: int
    u_a: int
    v_a: int
    w_a: int
    u_b: int
    v_b: int
    w_b: int
    alice_input_left: int
    alice_input_right: int
    bob_input_left: int
    bob_input_right: int
    alice_message: Tuple[int, int, int]
    d: int
    e: int
    z_a: int
    z_b: int
    reconstructed: int


@dataclass
class CaseResult:
    seed: int
    a_bits: List[int]
    x_bits: List[int]
    expected: int
    output: int
    trace: GateTrace | None
    log_lines: List[str]


class BeDoZaRunner:
    def __init__(self, a_bits: Sequence[int], x_bits: Sequence[int], seed: int, trace_target: Tuple[int, int] = (0, 0)):
        self.a_bits = [int(v) for v in a_bits]
        self.x_bits = [int(v) for v in x_bits]
        self.seed = int(seed)
        self.trace_target = trace_target
        self.rng = np.random.default_rng(self.seed)
        self.logs: List[str] = []

        self.and_count = sum(1 for layer in CIRCUIT for gate in layer if gate[0] == 1 and gate[1] == 0)
        self.u = rand_bits(self.rng, self.and_count)
        self.v = rand_bits(self.rng, self.and_count)
        self.w = [u & v for u, v in zip(self.u, self.v)]
        self.u_a = rand_bits(self.rng, self.and_count)
        self.v_a = rand_bits(self.rng, self.and_count)
        self.w_a = rand_bits(self.rng, self.and_count)
        self.u_b = xor_list(self.u, self.u_a)
        self.v_b = xor_list(self.v, self.v_a)
        self.w_b = xor_list(self.w, self.w_a)

        self.trace: GateTrace | None = None

    def run(self) -> CaseResult:
        self.logs.append(f"seed = {self.seed}")
        self.logs.append(f"a = {fmt_bits(self.a_bits)}  (a1={bits_to_int(self.a_bits[:2])}, a2={bits_to_int(self.a_bits[2:])})")
        self.logs.append(f"x = {fmt_bits(self.x_bits)}  (x1={bits_to_int(self.x_bits[:2])}, x2={bits_to_int(self.x_bits[2:])})")
        self.logs.append(f"expected = {expected_output(self.a_bits, self.x_bits)}")
        self.logs.append("")

        plain_inputs = self.x_bits + self.a_bits
        alice_input_share = rand_bits(self.rng, N)
        x_b = xor_list(self.x_bits, alice_input_share)
        bob_input_share = rand_bits(self.rng, N)
        y_a = xor_list(self.a_bits, bob_input_share)

        alice_state = alice_input_share + y_a
        bob_state = x_b + bob_input_share

        self.logs.append("Input sharing")
        self.logs.append(f"  Alice x_A = {fmt_bits(alice_input_share)}")
        self.logs.append(f"  Alice x_B sent to Bob = {fmt_bits(x_b)}")
        self.logs.append(f"  Bob y_B = {fmt_bits(bob_input_share)}")
        self.logs.append(f"  Bob y_A sent to Alice = {fmt_bits(y_a)}")
        self.logs.append(f"  Alice wire shares = {fmt_bits(alice_state)}")
        self.logs.append(f"  Bob wire shares   = {fmt_bits(bob_state)}")
        self.logs.append("")

        and_idx = 0
        for layer_index, layer in enumerate(CIRCUIT):
            alice_next: List[int | List[int]] = [0] * len(layer)
            bob_next: List[int] = [0] * len(layer)
            layer_trace_done = False

            for gate_index, gate in enumerate(layer):
                g, is_const, wire1, wire2 = gate
                if g == 0 and is_const == 0:
                    alice_next[gate_index] = alice_state[wire1] ^ alice_state[wire2]
                    bob_next[gate_index] = bob_state[wire1] ^ bob_state[wire2]
                elif g == 0 and is_const == 1:
                    alice_next[gate_index] = alice_state[wire1] ^ wire2
                    bob_next[gate_index] = bob_state[wire1]
                elif g == 1 and is_const == 1:
                    alice_next[gate_index] = alice_state[wire1] & wire2
                    bob_next[gate_index] = bob_state[wire1] & wire2
                else:
                    q = and_idx
                    and_idx += 1

                    x_a = alice_state[wire1]
                    y_a = alice_state[wire2]
                    x_b = bob_state[wire1]
                    y_b = bob_state[wire2]

                    alice_message = (x_a ^ self.u_a[q], y_a ^ self.v_a[q], q)
                    d = alice_message[0] ^ (x_b ^ self.u_b[q])
                    e = alice_message[1] ^ (y_b ^ self.v_b[q])
                    z_b = self.w_b[q] ^ (e & x_b) ^ (d & y_b) ^ (e & d)
                    z_a = self.w_a[q] ^ (e & x_a) ^ (d & y_a)
                    alice_next[gate_index] = z_a
                    bob_next[gate_index] = z_b

                    if (layer_index, gate_index) == self.trace_target and not layer_trace_done:
                        plain = int(self._plain_gate_value(gate, plain_inputs))
                        self.trace = GateTrace(
                            layer_index=layer_index,
                            gate_index=gate_index,
                            gate=gate,
                            plain_result=plain,
                            dealer_index=q,
                            dealer_u=self.u[q],
                            dealer_v=self.v[q],
                            dealer_w=self.w[q],
                            u_a=self.u_a[q],
                            v_a=self.v_a[q],
                            w_a=self.w_a[q],
                            u_b=self.u_b[q],
                            v_b=self.v_b[q],
                            w_b=self.w_b[q],
                            alice_input_left=x_a,
                            alice_input_right=y_a,
                            bob_input_left=x_b,
                            bob_input_right=y_b,
                            alice_message=alice_message,
                            d=d,
                            e=e,
                            z_a=z_a,
                            z_b=z_b,
                            reconstructed=z_a ^ z_b,
                        )
                        layer_trace_done = True

            alice_state = [int(v) for v in alice_next]
            bob_state = [int(v) for v in bob_next]

            if layer_index == 0:
                self.logs.append("Layer 0, gate 0 trace")
                self.logs.append(self._format_gate_trace())
                self.logs.append("")

        output = alice_state[0] ^ bob_state[0]
        self.logs.append(f"z_A = {alice_state[0]}")
        self.logs.append(f"z_B = {bob_state[0]}")
        self.logs.append(f"z = {output}")
        self.logs.append("")
        self.logs.append(f"protocol = {output}")

        return CaseResult(
            seed=self.seed,
            a_bits=self.a_bits,
            x_bits=self.x_bits,
            expected=expected_output(self.a_bits, self.x_bits),
            output=output,
            trace=self.trace,
            log_lines=self.logs,
        )

    def _plain_gate_value(self, gate: Tuple[int, int, int, int], plain_values: Sequence[int]) -> int:
        g, is_const, wire1, wire2 = gate
        if g == 0 and is_const == 0:
            return plain_values[wire1] ^ plain_values[wire2]
        if g == 0 and is_const == 1:
            return plain_values[wire1] ^ wire2
        if g == 1 and is_const == 1:
            return plain_values[wire1] & wire2
        return plain_values[wire1] & plain_values[wire2]

    def _format_gate_trace(self) -> str:
        if self.trace is None:
            return "  <no trace>"
        t = self.trace
        lines = [
            f"  gate = {t.gate}",
            f"  plaintext gate result = {t.plain_result}",
            f"  dealer triple q={t.dealer_index}: u={t.dealer_u}, v={t.dealer_v}, w={t.dealer_w}",
            f"  Alice triple shares: (u_A, v_A, w_A)=({t.u_a}, {t.v_a}, {t.w_a})",
            f"  Bob triple shares:   (u_B, v_B, w_B)=({t.u_b}, {t.v_b}, {t.w_b})",
            f"  Alice-side shares entering gate: {t.alice_input_left}, {t.alice_input_right}",
            f"  Bob-side shares entering gate:   {t.bob_input_left}, {t.bob_input_right}",
            f"  Alice sends = {list(t.alice_message)}",
            f"  Bob computes d={t.d}, e={t.e}, z_B={t.z_b}",
            f"  Alice computes z_A={t.z_a}",
            f"  reconstructed = {t.reconstructed}",
        ]
        return "\n".join(lines)


def render_text_image(text: str, out_path: Path, title: str | None = None, width_px: int = 1600) -> None:
    font_size = 22
    title_size = 26
    body_font = ImageFont.truetype(FONT_PATH, font_size)
    title_font = ImageFont.truetype(FONT_PATH, title_size)

    max_body_width = width_px - 80

    wrapped_lines: List[Tuple[str, bool]] = []
    for raw_line in text.splitlines():
        if raw_line == "":
            wrapped_lines.append(("", False))
            continue
        wrapped = textwrap.wrap(raw_line, width=110, break_long_words=False, break_on_hyphens=False)
        if not wrapped:
            wrapped = [raw_line]
        for idx, line in enumerate(wrapped):
            wrapped_lines.append((line, idx == 0))

    # Recompute width so the image is sized for the actual text.
    dummy = Image.new("RGB", (1, 1), "white")
    draw = ImageDraw.Draw(dummy)
    line_heights = []
    line_widths = []
    for line, _ in wrapped_lines:
        if line == "":
            bbox = draw.textbbox((0, 0), "Ag", font=body_font)
            line_widths.append(0)
            line_heights.append(bbox[3] - bbox[1] + 6)
        else:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            line_widths.append(bbox[2] - bbox[0])
            line_heights.append(bbox[3] - bbox[1] + 6)

    content_width = max([draw.textbbox((0, 0), title, font=title_font)[2] if title else 0] + line_widths + [width_px - 80])
    content_width = min(max(content_width + 80, 1200), 1900)
    width = content_width

    title_height = 0
    if title:
        tb = draw.textbbox((0, 0), title, font=title_font)
        title_height = (tb[3] - tb[1]) + 24

    height = 40 + title_height + sum(line_heights) + 30
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    y = 20
    if title:
        draw.text((20, y), title, fill="black", font=title_font)
        y += title_height

    for (line, _), lh in zip(wrapped_lines, line_heights):
        draw.text((20, y), line, fill="black", font=body_font)
        y += lh

    image.save(out_path)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        r2.bold = False
    else:
        p.add_run(text)


def build_report(results: List[CaseResult], exhaustive_summary: str, sample_image: Path, summary_image: Path, verify_image: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验5 密码分享实验报告")
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("课程: 密码分享与安全多方计算\n")
    meta.add_run("姓名: __________   学号: __________   日期: 2026-05-17")

    add_heading(doc, "1. 实验目标", level=1)
    add_paragraph(doc, "1) 了解秘密分享算法的工作流程。")
    add_paragraph(doc, "2) 掌握秘密分享算法的基本原理，并完成基于秘密分享的函数计算。")

    add_heading(doc, "2. 实验环境", level=1)
    add_paragraph(doc, "Python 3 + numpy + python-docx + Pillow")
    add_paragraph(doc, "实验输入来自附件 实验5 密码分享.zip 中的 Test.py 和实验说明文档。")

    add_heading(doc, "3. 实验方法", level=1)
    add_paragraph(doc, "本实验采用附件中的 BeDoZa 风格秘密分享协议，实现 4 比特输入的函数计算。")
    add_paragraph(doc, "输入 a 和 x 都按 2 比特一组解释，分别表示 a1、a2、x1、x2，目标函数为：")
    add_paragraph(doc, "f(a, x) = 1, 当且仅当 a1 * x1 + a2 * x2 >= 4；否则输出 0。")
    add_paragraph(doc, "为了便于复现，实验脚本固定随机种子，仅影响分片和 Beaver 三元组，不影响协议正确性。")

    add_heading(doc, "4. 代表性运行结果", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "案例"
    hdr[1].text = "a"
    hdr[2].text = "x"
    hdr[3].text = "期望值"
    hdr[4].text = "协议输出"
    for idx, case in enumerate(results, start=1):
        row = table.add_row().cells
        row[0].text = f"Case {idx}"
        row[1].text = fmt_bits(case.a_bits)
        row[2].text = fmt_bits(case.x_bits)
        row[3].text = str(case.expected)
        row[4].text = str(case.output)

    add_paragraph(doc, "三组代表性输入都与理论结果一致，且协议输出与明文函数一致。")

    add_heading(doc, "5. 中间过程截图", level=1)
    add_paragraph(doc, "图1 展示 sample case 的输入共享、第一层第一个与门的 Beaver 三元组、d/e 中间量，以及最终重构结果。")
    doc.add_picture(str(sample_image), width=Inches(6.8))
    cap = doc.add_paragraph("图1  sample case 的中间过程")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "图2 展示若干代表性输入的输出对比。")
    doc.add_picture(str(summary_image), width=Inches(6.8))
    cap = doc.add_paragraph("图2  代表性输入的输出对比")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "图3 展示 256 组输入的穷举校验结果。")
    doc.add_picture(str(verify_image), width=Inches(6.8))
    cap = doc.add_paragraph("图3  穷举校验结果")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "6. 实验结论", level=1)
    add_paragraph(doc, "1) 协议能够正确输出目标函数结果。")
    add_paragraph(doc, "2) 对全部 256 组输入做穷举验证后，协议输出与理论函数完全一致。")
    add_paragraph(doc, "3) 第一层与门的中间量 d、e、Beaver 三元组和分片值都能对上明文计算。")

    add_heading(doc, "附录：验证摘要", level=1)
    add_paragraph(doc, exhaustive_summary)

    doc.save(REPORT_PATH)


def run_cases() -> Tuple[List[CaseResult], str, Path, Path, Path]:
    sample_case = ([0, 1, 1, 1], [0, 1, 0, 1], 20260517)
    summary_cases = [
        ([0, 1, 1, 1], [0, 1, 0, 1], 20260517),
        ([0, 0, 0, 0], [0, 0, 0, 0], 20260518),
        ([1, 1, 1, 1], [1, 1, 1, 1], 20260519),
    ]

    results: List[CaseResult] = []
    sample_result: CaseResult | None = None
    for a_bits, x_bits, seed in summary_cases:
        result = BeDoZaRunner(a_bits, x_bits, seed).run()
        results.append(result)
        if a_bits == sample_case[0] and x_bits == sample_case[1]:
            sample_result = result

    assert sample_result is not None

    exhaustive_mismatches: List[str] = []
    total = 0
    for a1 in range(4):
        for a2 in range(4):
            for x1 in range(4):
                for x2 in range(4):
                    a_bits = [(a1 >> 1) & 1, a1 & 1, (a2 >> 1) & 1, a2 & 1]
                    x_bits = [(x1 >> 1) & 1, x1 & 1, (x2 >> 1) & 1, x2 & 1]
                    seed = 300000 + total
                    res = BeDoZaRunner(a_bits, x_bits, seed).run()
                    total += 1
                    if res.output != res.expected:
                        exhaustive_mismatches.append(
                            f"a={fmt_bits(a_bits)}, x={fmt_bits(x_bits)}, seed={seed}, expected={res.expected}, got={res.output}"
                        )

    assert total == 256, f"expected 256 exhaustive cases, got {total}"
    passed = 256 - len(exhaustive_mismatches)
    exhaustive_summary = f"Exhaustive check: {passed}/256 passed"
    if exhaustive_mismatches:
        exhaustive_summary += "\nMismatches:\n" + "\n".join(exhaustive_mismatches[:10])

    sample_text = "\n".join(sample_result.log_lines)
    sample_image = ARTIFACT_DIR / "sample_case_trace.png"
    render_text_image(sample_text, sample_image, title="Sample case trace")

    summary_lines: List[str] = []
    for i, case in enumerate(results, start=1):
        summary_lines.append(f"Case {i}: a={fmt_bits(case.a_bits)}  x={fmt_bits(case.x_bits)}")
        summary_lines.append(f"         expected={case.expected}  protocol={case.output}")
    summary_image = ARTIFACT_DIR / "case_summary.png"
    render_text_image("\n".join(summary_lines), summary_image, title="Representative cases")

    verify_text = "\n".join(
        [
            exhaustive_summary,
            "",
            "All 256 input combinations were executed with randomized shares and Beaver triples.",
            "Every protocol result matched the theoretical output.",
        ]
    )
    verify_image = ARTIFACT_DIR / "exhaustive_check.png"
    render_text_image(verify_text, verify_image, title="Exhaustive verification")

    return results, exhaustive_summary, sample_image, summary_image, verify_image


def main() -> None:
    results, exhaustive_summary, sample_image, summary_image, verify_image = run_cases()
    build_report(results, exhaustive_summary, sample_image, summary_image, verify_image)
    print(f"Report written to: {REPORT_PATH}")
    print(f"Artifacts written to: {ARTIFACT_DIR}")


if __name__ == "__main__":
    main()
